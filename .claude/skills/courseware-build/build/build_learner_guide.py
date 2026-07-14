#!/usr/bin/env python3
"""Generate the "Build a Human-AI Workforce with Autonomous AI Agents" Learner Guide as BOTH a Markdown
mirror (LG-*.md at repo root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
with commands · Test it), plus setup, wrap-up and glossary. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned with
the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# Pull the DETAILED step-by-step from each lab's README so the Learner Guide is
# highly detailed AND stays 100% aligned with the labs/ folder. Falls back to the
# terse course_data steps if a README is missing/unparseable.
import glob as _glob, re as _re
def _readme_steps(num):
    cands=_glob.glob(os.path.join(REPO,"labs",f"lab{num:02d}-*","README.md"))
    if not cands: return None
    txt=open(cands[0],encoding="utf-8").read()
    # Terminate the Steps body only at a 2-hash section heading ("## Verification"),
    # NOT at 3-hash "### N." step headers used by some labs.
    m=_re.search(r'\n##\s*Steps\s*\n(.*?)(?=\n##\s|\Z)', txt, _re.S)
    if not m: return None
    body=m.group(1).strip()
    # Handle both step formats: "N. **Title.** prose" list items AND "### N. Title" headers.
    items=_re.split(r'\n(?=(?:#{2,4}\s*)?\d+\.\s)', body)
    out=[]
    for it in items:
        it=it.strip()
        if not _re.match(r'(?:#{2,4}\s*)?\d+\.\s', it): continue
        cm=_re.search(r'```[a-zA-Z0-9]*\n(.*?)```', it, _re.S)
        cmd=cm.group(1).strip() if cm else ""
        # ignore non-command fenced snippets (e.g. ```text chat prompts) — keep only real commands
        if cmd and _re.match(r'^\s*#', cmd) and "\n" not in cmd:
            cmd=""
        text=it[:cm.start()] if cm else it
        # drop markdown table rows (| ... |) — tables don't flatten into prose
        text="\n".join(l for l in text.splitlines() if not l.strip().startswith("|"))
        text=_re.sub(r'^(?:#{2,4}\s*)?\d+\.\s*','',text)      # drop leading '### N.' / 'N.'
        text=_re.sub(r'\*\*(.*?)\*\*', r'\1', text)           # drop bold markers
        text=_re.sub(r'\*(.*?)\*', r'\1', text)               # drop italics
        text=_re.sub(r'`([^`]+)`', r'\1', text)               # drop inline code ticks
        text=_re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)  # link -> text (url)
        text=_re.sub(r'\s*\n\s*',' ', text).strip()
        if text:
            out.append((text,cmd))
    return out or None

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It provides step-by-step instructions for all 34 hands-on labs across three autonomous-AI-agent "
  "platforms — Hermes Agent, OpenClaw and Paperclip — organised into three topics. Each platform builds "
  "on the same skillset: install the runtime, connect models and channels, give the agent memory, extend "
  "it with skills and tools, automate it with schedules, secure it with governance, and finally "
  "orchestrate multiple agents to deliver a realistic business outcome.")
p("Use this guide alongside the course slides and the lab files in the labs/ folder of the course "
  "repository. The labs run on your own laptop and in Docker; where a lab connects to an external "
  "service (an LLM provider, a messaging channel, a VPS) use only accounts and credentials you own, "
  "and keep API keys and tokens out of prompts and out of version control.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Before You Start — Environment Setup")
h3("What you need")
bullets([
 "A laptop you have administrator rights on — Windows 10/11, macOS 12+ or Ubuntu 22.04+ — with a terminal.",
 "Node.js 24 LTS (required to run the OpenClaw gateway daemon).",
 "Docker Desktop (required to self-host Paperclip with Docker Compose, and used for isolated agent sandboxes).",
 "At least one LLM provider login or API key — Claude (Anthropic), OpenAI, OpenRouter, MiniMax or DeepSeek.",
 "A free Telegram account for the messaging-channel labs (you will create a bot via BotFather).",
 "Optional: a VPS (e.g. Hostinger or exe.dev) if you want to run an agent 24/7 beyond the classroom.",
])
h3("Verify your toolchain")
p("Open a terminal and confirm the core prerequisites are present before you begin. Each platform's "
  "installer (Hermes, OpenClaw, Paperclip) is run in its own lab; here you only confirm the building blocks.")
code("$ node --version     # Node.js 24.x for OpenClaw\n$ docker --version   # Docker Desktop for Paperclip and sandboxes\n$ docker compose version\n$ git --version      # to clone the course labs repository")
h3("Conventions used in every lab")
bullets([
 "Commands are run from your terminal; a leading sudo is used only where elevated rights are required.",
 "Placeholders such as <API_KEY>, <BOT_TOKEN> and <HOST> are replaced with your own values.",
 "Store provider keys and channel tokens in each platform's config (never in prompts or in git).",
 "Keep approval prompts, budgets and sandboxing on so agents act under human oversight.",
 "Each capstone orchestrates multiple agents — start the single-agent labs first so the pieces are familiar.",
])

# ---------------- per-topic, per-lab ----------------
TOPICS_BY_NUM={t["num"]:t for t in C.TOPICS}
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets(t["concepts"])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Learning outcome: {a['objective']}.")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        h3("Step-by-step")
        # Prefer the detailed README steps. The LG never shows YouTube reference
        # links — videos live in the labs/ READMEs only.
        st=_readme_steps(a["num"]) or [(instr,cmd) for (instr,cmd) in a["steps"]]
        st=[(si,sc) for (si,sc) in st if "youtube" not in sc.lower() and "youtube" not in si.lower()]
        steps(st)
        h3("Test it")
        p(a["test"])
        note(f"Full commands and screenshots are in labs/lab-{a['num']:02d}-*.md. "
             f"Use only accounts, keys and hosts you own, and keep agents under human oversight.")
        rule()

h1("Wrap-Up — The Agent Build Arc and Governance")
p("These cross-cutting themes run through every platform in the course. Study this section alongside the labs "
  "so the same skillset transfers from Hermes to OpenClaw to Paperclip and into your own production agents.")
h3("The agent build arc")
p("Every platform follows the same progression — the order in which you built each agent in the labs:")
bullets([
 "Install — set up the runtime (Hermes CLI/TUI, the OpenClaw Node.js gateway daemon, or Paperclip via Docker Compose).",
 "Models — connect an LLM provider (Claude, OpenAI, OpenRouter, MiniMax or DeepSeek) and pick a capable model.",
 "Channels — reach the agent where work happens (Telegram, WhatsApp, Discord, Slack and more).",
 "Memory — give the agent durable, user- or company-scoped recall across sessions and surfaces.",
 "Skills & tools — extend the agent with reusable skills and third-party tool/integration calls.",
 "Crons & heartbeat — schedule recurring work and let a heartbeat keep the agent running.",
 "Security — isolate execution in sandboxes, store secrets in config, and gate risky actions with approvals.",
 "Multi-agent — orchestrate a coordinator plus specialist workers to deliver an end-to-end business outcome.",
])
h3("Human-in-the-loop governance")
bullets([
 "Approvals — important or destructive actions pause for explicit human sign-off before they run.",
 "Budgets — company- and per-agent spend caps warn at 80% and pause at 100% so costs never run away.",
 "Sandboxing — agents execute in isolated backends (Docker, SSH, or a remote host) so the host stays safe.",
 "Audit — an activity log and cost-events trail record who did what, giving you a CFO-style view of the workforce.",
 "Least privilege — scope each channel, skill and tool with allow/deny lists so agents can only do their job.",
])
rule()

h1("Next Steps")
bullets([
 "First pass: complete every lab on your own machine, following the commands in each lab file.",
 "Second pass: redo the labs from memory until installing, connecting and securing an agent is automatic.",
 "Extend the labs with your own skills, tools and integrations for a workflow you care about.",
 "Deploy an agent on a VPS or cloud backend for 24/7 operation, with budgets, approvals and audit turned on.",
 "Explore multi-agent orchestration — a coordinator plus specialists — for a realistic end-to-end business task.",
 "Review each lab's detailed steps in this guide and re-run the labs on your own machine.",
])

h1("Glossary")
gl=[
 ("Autonomous AI agent","An LLM-driven system that perceives, reasons, acts with tools and remembers, running under human oversight."),
 ("Agent stack","The layers that make an agent work: model, memory, skills, tools, channels and a scheduler/security layer."),
 ("Model / provider","The LLM (and the vendor serving it) that gives the agent its reasoning — e.g. Claude, OpenAI, OpenRouter."),
 ("Channel","A messaging surface the agent is reachable on, such as Telegram, WhatsApp, Discord or Slack."),
 ("Memory","Durable, user- or company-scoped storage that lets the agent recall context across sessions and surfaces."),
 ("Skill","A reusable, named capability an agent can install or create and invoke on demand."),
 ("Tool / integration","An external action or service the agent can call — web search, image generation, email, GitHub via MCP."),
 ("MCP","Model Context Protocol — a standard way to connect an agent to external tool servers."),
 ("Cron / heartbeat","A schedule that triggers recurring agent work; a heartbeat self-checks and keeps the agent running."),
 ("Sandbox","An isolated execution backend (Docker, SSH, remote host) that contains what an agent's commands can touch."),
 ("Approval gate","A human sign-off required before an agent performs an important or destructive action."),
 ("Budget cap","A spend limit at company or per-agent level that warns at 80% and pauses work at 100%."),
 ("Multi-agent orchestration","A coordinator agent delegating to specialist worker agents to deliver an end-to-end outcome."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("1.0","12 July 2026","Initial release — Learner Guide covering the autonomous-AI-agent labs (Hermes Agent, OpenClaw, Paperclip).",C.TRAINER),
 ("1.1","14 July 2026","Revised for the 2-day structure and 34 labs (Hermes 14, OpenClaw 12, Paperclip 8) with per-lab reference videos; learning outcomes aligned to the Course Proposal / TSC.",C.TRAINER),
 ("1.2",C.VERSION_DATE,"Per-lab step numbering; YouTube links removed (videos remain in the labs); Hermes install lab expanded with the official install methods.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        for i,(instr,cmd) in enumerate(rest[0],1):
            # Literal per-lab step numbers (Word's List Number style numbers
            # continuously across the whole document — 1..N must restart per lab).
            para=doc.add_paragraph()
            para.paragraph_format.left_indent=Pt(18)
            para.paragraph_format.space_after=Pt(4)
            r=para.add_run(f"{i}.  "); r.bold=True
            para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
