#!/usr/bin/env python3
"""Generate the "Build a Human-AI Workforce with Autonomous AI Agents" Lesson Plan (LP) DOCX in the Tertiary house format.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 3-day schedule tables (9:30am-6:30pm, 8 training hours/day, 1h
lunch, tea within, final assessment Day 3). Day themes, topics and labs come
from course_data + the domain data files so the LP stays aligned with the deck,
guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def lab_titles(nums):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in ACT if a['num'] in nums)

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, activity_text)  kind: admin/topic/lab/break/lunch/assess/recap
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30","9:50",20,"admin","Welcome, course introduction, ground rules and mandatory digital attendance (AM)"),
    ("9:50","10:30",40,"topic","Topic 1 — "+C.TOPICS[0]["title"]+": what an autonomous AI agent is, the agent stack (model, memory, skills, tools) and human-in-the-loop governance (concepts + demo)"),
    ("10:30","11:00",30,"lab","Hands-on: "+lab_titles([1,2,3])),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:00",105,"lab","Hands-on: "+lab_titles([4,5,6,7,8,9])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","14:40",40,"lab","Hands-on: "+lab_titles([10,11,12,13,14])),
    ("14:40","15:20",40,"topic","Topic 2 — "+C.TOPICS[1]["title"]+": the OpenClaw gateway daemon, providers, channels, skills, tools, dreaming and crons for a business back-office (concepts + demo)"),
    ("15:20","15:35",15,"break","Tea break"),
    ("15:35","18:15",160,"lab","Hands-on: "+lab_titles([15,16,17,18,19,20,21,22,23,24,25,26])),
    ("18:15","18:30",15,"recap","Day 1 recap, Q&A and PM digital attendance"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:30","9:45",15,"recap","Day 1 recap and mandatory digital attendance (AM)"),
    ("9:45","10:30",45,"topic","Topic 3 — "+C.TOPICS[2]["title"]+": Docker-hosted company of agents, connecting OpenAI Codex, human approval gates, task tracking, budgets and audit (concepts + demo)"),
    ("10:30","11:00",30,"lab","Hands-on: "+lab_titles([27,28])),
    ("11:00","11:15",15,"break","Tea break"),
    ("11:15","13:00",105,"lab","Hands-on: "+lab_titles([29,30,31,32])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","15:30",90,"lab","Hands-on: "+lab_titles([33,34])+". Course recap and next steps"),
    ("15:30","15:45",15,"break","Tea break"),
    ("15:45","16:15",30,"recap","Consolidation, Q&A and assessment preparation"),
    ("16:15","16:30",15,"assess","Briefing for Assessment"),
    ("16:30","17:30",60,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, open book"),
    ("17:30","18:30",60,"assess","Practical Performance (PP) — hands-on agent-building tasks that follow the course labs, 1 hour, open book. PM digital attendance"),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("1.0","12 July 2026","Initial release — 3-day lesson plan for the autonomous-AI-agent labs (Hermes Agent, OpenClaw, Paperclip).",C.TRAINER),
 ("1.1","14 July 2026","Revised to a 2-day plan (Day 1: Topics 1 & 2; Day 2: Topic 3 & final assessment). Hermes expanded to 14 labs, OpenClaw to 12, Paperclip to 8 (34 labs total). Learning outcomes aligned to the Course Proposal / TSC; per-lab reference videos added.",C.TRAINER),
 ("1.2","14 July 2026","Visual redesign: all content slides card-format; YouTube links removed from the deck and Learner Guide (videos remain in the labs); live Hermes screenshots and skill-category grids added; Hermes install lab expanded with the official install methods.",C.TRAINER),
 ("1.3","14 July 2026","Deck restyled in the dark masterclass theme end-to-end; added the Hermes masterclass slides (What is Hermes; The Bet: Compounding Value; platform matrix; Top-10 slash commands; Telegram setup; memory L1-vs-L2 and session-search; what-is-a-tool, built-in tool catalog, tools-in-action, cron scheduling, delegation, kanban orchestration and OpenClaw topic slides); Lab 13 rebuilt as the multi-agent video team (profiles + kanban) with detailed Labs 11-14 guides; fade slide transitions and the live Achievements dashboard screenshot.",C.TRAINER),
 ("1.4","15 July 2026","Security masterclass slides added to the Hermes Security lab (defense-in-depth layers; the gateway's six-check default-deny trust chain; channel lockdown dials; the three approval modes + the Tirith scanner; a live Dangerous-Command approval; YOLO mode; the unrecoverable-command blocklist floor; prompt-injection scanning of context files). Security lab steps expanded with the channel allowlist and approvals.mode.",C.TRAINER),
 ("1.5",C.VERSION_DATE,"Sixteen concept diagrams imported from the original Mastering OpenClaw masterclass deck into Topic 2 (OpenClaw OS control plane; deployment; hub-and-spoke gateway; tools vs skills; SKILL-on-demand; 6-phase execution loop; exec; HEARTBEAT + autonomic tick; workspace file persistence; context assembly; compaction; defense methods; poisoned-skill anatomy; sessions_spawn; sub-agent routing).",C.TRAINER),
])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration","2 days · 8 training hours per day (16 hours)"),
      ("Daily Timing","9:30 am – 6:30 pm (1-hour lunch; tea breaks within training time)"),
      ("Mode","Instructor-led, hands-on autonomous-AI-agent labs (local machine + Docker)"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(10.5)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Final assessment is conducted on Day 2 from 4:30 pm.",C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(10.5)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    training=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        if kind!="lunch": training+=mins
    # widths
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    p=doc.add_paragraph(); r=p.add_run(f"Total training time: {training} minutes ({training//60} hours)."); r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert training==480, f"Day {day} training minutes = {training}, expected 480"

H("Lab Reference (aligned to course topics)",1)
tt=doc.add_table(rows=0,cols=3); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic / Platform","Weighting","Labs"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],tp["weighting"],size=9.5,fill=TOPIC_FILL)
    set_cell(cells[2],", ".join(f"Lab {a['num']}" for a in acts),size=9.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
