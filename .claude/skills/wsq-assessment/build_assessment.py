#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Application Integration with Docker and Kubernetes' (TGS-2021010366):
  - Written Assessment (SAQ)  — 5 open-ended KNOWLEDGE questions (K1–K5), aligned to the slides
  - Practical Performance (PP) — 4 PRACTICAL tasks (LO1–LO4), aligned to the in-class activities
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Build a Human-AI Workforce with Autonomous AI Agents"
COURSE_CODE = "TGS-2024043854"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("course-logo.png")   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v3", "v3"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
WRITTEN = [
 ("K1",
  "A Singapore logistics company is preparing an AI adoption roadmap and is mapping the landscape of AI "
  "applications - including autonomous AI agents such as Hermes Agent, OpenClaw and Paperclip - before deciding "
  "where to deploy them.",
  "Describe three (3) distinct categories of AI application used across different industries and state one (1) key "
  "capability associated with each category.",
  ["Large Language Model (LLM)-based applications, such as AI chatbots and text-generation tools - key capability: "
   "understand and produce human-like natural language for automated engagement and content generation.",
   "Computer vision applications, such as defect detection and visual product search - key capability: analyse "
   "images and video to automate visual inspection and classification.",
   "Predictive analytics / machine-learning applications, such as demand forecasting and recommendation engines - "
   "key capability: model historical data to forecast outcomes and support data-driven decisions.",
   "Autonomous AI agent applications, such as Hermes Agent, OpenClaw and Paperclip - key capability: execute "
   "multi-step workflows (research, scheduling, orchestration) with memory, skills and tools and minimal human "
   "intervention. (Slides: Autonomous AI Agents Fundamentals / Three Platforms, One Skillset)"]),
 ("K2",
  "A Singapore healthcare provider has deployed a Retrieval Augmented Generation (RAG) agent that also uses "
  "persistent memory to help its care team query internal clinical guidelines. Clinicians report inconsistent "
  "answer quality, and the product manager must review performance effectiveness before recommending fixes.",
  "Explain the concept of 'performance effectiveness' in the context of a Retrieval Augmented Generation (RAG) "
  "agent and describe two (2) factors - including how the agent's memory or retrieval is configured - that may "
  "contribute to inconsistent retrieval and generation quality.",
  ["Performance effectiveness is the degree to which the agent retrieves contextually relevant documents from its "
   "knowledge base/memory and generates responses that are factually grounded, coherent and aligned with the "
   "query intent.",
   "Embedding / retrieval quality: a mismatch between the query and the stored document/memory chunks returns "
   "irrelevant or incomplete context, degrading answer accuracy.",
   "Chunk size and overlap, and how much conversation memory is retained: too little context omits critical "
   "information while too much adds noise, both reducing relevance.",
   "Source-document / memory quality: outdated or ambiguous content in the knowledge base or agent memory "
   "increases inaccurate or contradictory responses. (Slides: Memory & Plugins / MCP and Tools - RAG)"]),
 ("K3",
  "A Singapore bank has integrated an AI agent into its loan-processing product. The head of product is running a "
  "post-deployment review to check whether the agent meets its intended performance objectives.",
  "Describe three (3) methods that an organisation can apply to evaluate the effectiveness of an AI agent "
  "application following deployment and state the primary metric associated with each method.",
  ["Accuracy evaluation against a validated test dataset - primary metric: accuracy rate or F1-score.",
   "User acceptance testing under real workflow conditions - primary metric: user-satisfaction score or "
   "task-completion rate.",
   "Latency and throughput testing under load - primary metric: average response time (ms) and sustained "
   "concurrent requests.",
   "Business-impact assessment against a pre-deployment baseline - primary metric: percentage change in a target "
   "outcome such as processing time, error rate or cost per transaction. (Slides: Security / Paperclip audit - "
   "activity_log and cost_events)"]),
 ("K4",
  "A Singapore telecommunications company is building an agentic AI support agent using skills, tools, subagents "
  "and memory. The development team is designing the agent's underlying algorithm and needs guidance on how "
  "algorithm and agent design affect expected performance.",
  "Explain how the design of an algorithm (agent design) influences the performance of an agentic AI chatbot and "
  "describe two (2) design considerations - drawn from prompt design, skills/tools orchestration, subagent "
  "delegation or memory/context management - that the team should address during implementation.",
  ["Agent design determines the sequence of operations the agent follows to interpret input, recall memory, "
   "select skills/tools and construct a response; a well-structured design ensures each step produces "
   "sufficient-quality output for accurate interactions.",
   "Prompt engineering: the structure, specificity and scope of the system prompt directly influence response "
   "quality; poor prompts produce ambiguous or off-topic output.",
   "Skills and tool/MCP orchestration: which skills and tools the agent may invoke and under what conditions; weak "
   "orchestration causes wrong-tool selection, redundant actions or loops.",
   "Subagent delegation and context/memory management: delegating sub-tasks to worker subagents and keeping "
   "conversation history, retrieved documents and memory within the model's token limit so earlier context is not "
   "truncated. (Slides: Skills / MCP and Tools / Subagents & Delegation / Memory & Plugins)"]),
 ("K5",
  "A Singapore SaaS company has incorporated AI agents (with crons and automations) into its product-management "
  "workflow. The product director is running a quarterly review to assess whether the AI integration produced "
  "measurable improvements.",
  "Describe two (2) methods that an organisation can apply to evaluate process improvements to product management "
  "following AI-agent integration and identify one (1) key indicator of improvement for each method.",
  ["Baseline comparison analysis - measure the process before and after AI-agent integration on consistent "
   "metrics; key indicator: percentage reduction in cycle time for targeted tasks.",
   "Controlled pilot deployment - run the AI-assisted workflow against the original process within a defined "
   "scope; key indicator: difference in output accuracy or error rate versus the control group.",
   "Stakeholder feedback review - collect structured input on utility and reliability; key indicator: acceptance "
   "rate of AI-generated recommendations without manual correction.",
   "Return-on-investment analysis - quantify value against cost; key indicator: cost saving per unit of output or "
   "revenue uplift. (Slides: Security / Budgets / Cron Jobs & Automation)"]),
 ("K6",
  "A newly established Singapore startup is building its product-development framework and is evaluating how AI "
  "agents and agentic workflows can be applied at different stages of its product lifecycle.",
  "Identify three (3) stages of the product development track where AI agent applications can be applied and "
  "explain one (1) specific use case for each stage.",
  ["Ideation and requirements gathering - an agent analyses customer feedback and market research to surface "
   "recurring themes and unmet needs.",
   "Design and prototyping - a multi-agent workflow drafts user stories, specifications and interface concepts, "
   "reducing time-to-first-deliverable.",
   "Testing and quality assurance - an agent generates test cases, detects anomalies and predicts high-risk defect "
   "areas, improving coverage.",
   "Post-launch monitoring and maintenance - autonomous agents (e.g. a Paperclip company of agents) monitor "
   "metrics, summarise incidents and recommend corrective actions. (Slides: Multi-Agent Workflow / Track AI Tasks "
   "/ Assign Task)"]),
]

SCENARIO = (
 "You are an automation engineer at a Singapore technology consultancy. Across this course you built and operated "
 "autonomous AI agents on three platforms - Hermes Agent (Topic 1), OpenClaw (Topic 2) and Paperclip (Topic 3). "
 "Complete the three tasks below; each maps to one platform and mirrors the hands-on labs you did in class. For "
 "EACH task, carry out the implementation on the platform and PASTE A SCREENSHOT of your working implementation "
 "as evidence, then write the required analysis.")

# (label, criterion, task prompt, box caption, model-answer build steps citing the activity)
BOX_CAP = "Paste a SCREENSHOT of your implementation, plus your written analysis, in the box below"
PRACTICAL = [
 ("Task 1", "A1, A3",
  "HERMES AGENT (Topic 1). A logistics client wants an autonomous LLM-based agent for shipment tracking and "
  "customer enquiries. Using Hermes Agent as you did in the labs (Lab 1 Install, Lab 5 Providers & Model, Lab 3 "
  "Memory, Lab 4 Skills, Lab 6 MCP & Tools), install and configure an agent connected to an LLM provider, then "
  "analyse the algorithm underlying the LLM-based application and identify two (2) strengths and two (2) "
  "limitations of such an application for the logistics client's operations. Paste a screenshot of your running "
  "Hermes agent (e.g. 'hermes doctor' / the TUI with your model set).",
  BOX_CAP,
  "Build steps (Hermes, from the labs): install ('curl -fsSL https://hermes-agent.nousresearch.com/install.sh | "
  "bash'), setup ('hermes setup --portal'), select a model ('hermes model'), add memory/skills/tools ('hermes "
  "skills install ...', MCP in ~/.hermes/config.yaml). Screenshot evidence: 'hermes doctor' green or the TUI. "
  "(Labs 1, 3, 4, 5, 6.)\n"
  "Analysis of the LLM algorithm: transformer-based network predicting token sequences from patterns learned in "
  "training, handling diverse open-ended input without task-specific rules.\n"
  "Strength 1 - interprets unstructured, varied customer enquiries without predefined decision trees.\n"
  "Strength 2 - cross-task adaptability: one agent handles tracking, escalation and notification drafting.\n"
  "Limitation 1 - hallucination: plausible but factually wrong output, a risk where responses must reflect live "
  "shipment data.\n"
  "Limitation 2 - dependence on training-data currency: without live tools/retrieval, answers may not reflect "
  "real-time status."),
 ("Task 2", "A2, A6",
  "OPENCLAW (Topic 2). An e-commerce client runs an agentic AI chatbot for pre-sales enquiries and wants better "
  "response efficiency. Using OpenClaw as you did in the labs (Lab 15 Install, Lab 17 Channels, Lab 18 Skills, Lab "
  "19 Tools, Lab 22 Memory), build an agentic chatbot on a channel with skills/tools, then establish the "
  "correlation between your agent's design (prompt, skills/tools orchestration, memory) and its operational "
  "efficiency, and assess two (2) AI-driven improvements to enhance its performance in product management. Paste a "
  "screenshot of your OpenClaw agent replying on a channel (e.g. Telegram).",
  BOX_CAP,
  "Build steps (OpenClaw, from the labs): install ('npm install -g openclaw@latest', 'openclaw onboard'), add a "
  "channel ('openclaw channel add telegram --token <TOKEN>', 'openclaw channel start telegram'), add a skill "
  "('openclaw skills add <name>'), connect a tool and enable memory. Screenshot evidence: the bot replying in the "
  "channel. (Labs 15, 17, 18, 19, 22.)\n"
  "Correlation: prompt architecture, context/memory configuration and skill/tool orchestration drive accuracy, "
  "relevance and latency; shallow context or weak prompts produce off-topic answers needing human follow-up, "
  "lowering efficiency; good memory + structured retrieval resolves enquiries in fewer turns.\n"
  "Improvement 1 - integrate a product knowledge base via a RAG tool so the agent surfaces accurate, current "
  "specifications instead of static training data.\n"
  "Improvement 2 - add a feedback loop that flags low-confidence responses for human review and uses validated "
  "corrections (and memory) to refine behaviour over time."),
 ("Task 3", "A4, A5",
  "PAPERCLIP (Topic 3). A financial-services client is choosing between a retrieval augmented generation (RAG) "
  "approach and a standard LLM assistant to run its client-advisory product as a company of AI agents. Using "
  "Paperclip as you did in the labs (Lab 27 Setup, Lab 28 Connect OpenAI Codex, Lab 31 Automate Hiring, Lab 32 "
  "Setup Agents), set up the company of agents, then evaluate the RAG approach versus the standard LLM assistant "
  "(strengths and limitations) and assess the feasibility of implementing the RAG approach within the client's "
  "product and its maintenance processes. Paste a screenshot of your Paperclip dashboard (company + agents).",
  BOX_CAP,
  "Build steps (Paperclip, from the labs): self-host ('git clone https://github.com/paperclipai/paperclip.git', "
  "'docker compose -f docker-compose.quickstart.yml up --build', open http://localhost:3100), connect the OpenAI "
  "Codex adapter, hire/approve agents, set up agents with mandates and budgets. Screenshot evidence: the dashboard "
  "showing the company and active agents. (Labs 27, 28, 31, 32.)\n"
  "RAG strengths: retrieves from a curated, firm-maintained corpus before generating, reducing hallucination and "
  "improving auditable factual accuracy - important in regulated advisory contexts.\n"
  "RAG limitation: operational overhead of maintaining and indexing the corpus; outdated or poorly indexed "
  "documents degrade retrieval quality.\n"
  "Standard LLM assistant strengths: faster to deploy, no document pipeline; limitation: no access to current "
  "firm data and no auditable response provenance (a regulatory concern).\n"
  "Feasibility: RAG is more suitable for client advisory given accurate, retrievable, auditable answers; "
  "implementation must budget for ongoing corpus curation, access-control management and version governance, "
  "governed by Paperclip's approval gates and budget rails as core maintenance obligations."),
]

def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    # Pagination is EXPLICIT — two questions to a page on the paper, one model answer to a
    # page in the key. Do not swap this for Word's keepNext/cantSplit: Word pushes an
    # oversized box to the next page, but Google Docs draws the border anyway and prints the
    # question text and the page footer straight THROUGH it. See SKILL.md → Pagination.
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    # Practical tasks are long and their boxes are tall, so they get a page each — on the
    # paper AND in the key. Same rule as the WA: the page break is ours, not the renderer's.
    for i, (label, crit, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
        if i < len(PRACTICAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
